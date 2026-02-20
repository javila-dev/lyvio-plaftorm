import os
import logging
from pathlib import Path
from typing import Dict
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentAnalyzer:
    """Servicio para analizar y validar documentos antes de procesarlos"""
    
    # Límites por tipo de documento
    MAX_PAGES = {
        '.pdf': 30,      # Máximo 30 páginas para PDF
        '.doc': 25,      # Máximo 25 páginas para Word
        '.docx': 25,
        '.txt': None,    # Para txt validamos por tamaño de texto
        '.md': None,
    }
    
    # Límites de caracteres por tipo
    MAX_CHARS = {
        '.txt': 50000,   # ~50k caracteres para txt (~10k palabras)
        '.md': 50000,
    }
    
    # Tamaño máximo de archivo en MB (debe coincidir con Document.MAX_FILE_SIZE_MB)
    MAX_FILE_SIZE_MB = 10
    
    # Mínimo de texto útil (en caracteres)
    MIN_TEXT_LENGTH = 100
    
    # Ratio máximo de imágenes a texto para PDFs
    MAX_IMAGE_TO_TEXT_RATIO = 0.5  # 50% del contenido puede ser imágenes
    
    @staticmethod
    def analyze_document(file_obj, filename: str) -> Dict:
        """
        Analiza un documento y retorna información sobre su contenido
        
        Returns:
            dict con:
            - is_valid: bool
            - error: str (si is_valid=False)
            - stats: dict con estadísticas del documento
        """
        file_ext = Path(filename).suffix.lower()
        file_size_mb = file_obj.size / (1024 * 1024)
        
        result = {
            'is_valid': True,
            'error': None,
            'stats': {
                'filename': filename,
                'file_type': file_ext,
                'file_size_mb': round(file_size_mb, 2),
                'pages': 0,
                'text_length': 0,
                'word_count': 0,
                'has_images': False,
                'image_pages': 0,
            }
        }
        
        # Validar tamaño de archivo
        if file_size_mb > DocumentAnalyzer.MAX_FILE_SIZE_MB:
            result['is_valid'] = False
            result['error'] = f"El archivo excede el tamaño máximo de {DocumentAnalyzer.MAX_FILE_SIZE_MB}MB"
            return result
        
        try:
            if file_ext == '.pdf':
                return DocumentAnalyzer._analyze_pdf(file_obj, result)
            elif file_ext in ['.doc', '.docx']:
                return DocumentAnalyzer._analyze_docx(file_obj, result)
            elif file_ext in ['.txt', '.md']:
                return DocumentAnalyzer._analyze_text(file_obj, result, file_ext)
            else:
                result['is_valid'] = False
                result['error'] = f"Tipo de archivo no soportado: {file_ext}"
                return result
                
        except Exception as e:
            logger.error(f"Error analizando documento {filename}: {str(e)}")
            result['is_valid'] = False
            result['error'] = f"Error al analizar el documento: {str(e)}"
            return result
    
    @staticmethod
    def _analyze_pdf(file_obj, result: Dict) -> Dict:
        """Analiza un archivo PDF"""
        try:
            file_obj.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_obj)
            num_pages = len(pdf_reader.pages)
            
            result['stats']['pages'] = num_pages
            
            # Validar número de páginas
            max_pages = DocumentAnalyzer.MAX_PAGES['.pdf']
            if num_pages > max_pages:
                result['is_valid'] = False
                result['error'] = f"El PDF tiene {num_pages} páginas. Máximo permitido: {max_pages} páginas"
                return result
            
            # Extraer texto de todas las páginas
            text = ""
            image_pages = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    text += page_text
                    
                    # Detectar páginas con imágenes (poco texto)
                    if len(page_text.strip()) < 50:  # Menos de 50 caracteres = probablemente imagen
                        image_pages += 1
                        
                except Exception as e:
                    logger.warning(f"Error extrayendo texto de página {page_num}: {str(e)}")
                    continue
            
            result['stats']['text_length'] = len(text)
            result['stats']['word_count'] = len(text.split())
            result['stats']['image_pages'] = image_pages
            result['stats']['has_images'] = image_pages > 0
            
            # Validar ratio de imágenes a texto
            if num_pages > 0:
                image_ratio = image_pages / num_pages
                if image_ratio > DocumentAnalyzer.MAX_IMAGE_TO_TEXT_RATIO:
                    result['is_valid'] = False
                    result['error'] = f"El PDF tiene demasiadas imágenes ({image_pages}/{num_pages} páginas). Máximo permitido: {int(DocumentAnalyzer.MAX_IMAGE_TO_TEXT_RATIO * 100)}% del contenido"
                    return result
            
            # Validar que tenga texto útil
            if len(text.strip()) < DocumentAnalyzer.MIN_TEXT_LENGTH:
                result['is_valid'] = False
                result['error'] = "El PDF no contiene suficiente texto útil para procesar"
                return result
            
            logger.info(f"PDF analizado: {num_pages} páginas, {len(text)} caracteres, {image_pages} páginas con imágenes")
            
        except Exception as e:
            result['is_valid'] = False
            result['error'] = f"Error al leer el PDF: {str(e)}"
        
        return result
    
    @staticmethod
    def _analyze_docx(file_obj, result: Dict) -> Dict:
        """Analiza un archivo Word (.docx)"""
        try:
            file_obj.seek(0)
            doc = DocxDocument(file_obj)
            
            # Contar párrafos como "páginas" aproximadas (4-5 párrafos = 1 página)
            num_paragraphs = len(doc.paragraphs)
            estimated_pages = max(1, num_paragraphs // 4)
            
            result['stats']['pages'] = estimated_pages
            
            # Validar número de páginas estimadas
            max_pages = DocumentAnalyzer.MAX_PAGES['.docx']
            if estimated_pages > max_pages:
                result['is_valid'] = False
                result['error'] = f"El documento tiene aproximadamente {estimated_pages} páginas. Máximo permitido: {max_pages} páginas"
                return result
            
            # Extraer todo el texto
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            result['stats']['text_length'] = len(text)
            result['stats']['word_count'] = len(text.split())
            
            # Detectar imágenes
            has_images = any(rel.target.endswith(('.png', '.jpg', '.jpeg', '.gif')) 
                           for rel in doc.part.rels.values() 
                           if hasattr(rel, 'target'))
            result['stats']['has_images'] = has_images
            
            # Validar que tenga texto útil
            if len(text.strip()) < DocumentAnalyzer.MIN_TEXT_LENGTH:
                result['is_valid'] = False
                result['error'] = "El documento no contiene suficiente texto útil para procesar"
                return result
            
            logger.info(f"DOCX analizado: ~{estimated_pages} páginas, {len(text)} caracteres")
            
        except Exception as e:
            result['is_valid'] = False
            result['error'] = f"Error al leer el documento Word: {str(e)}"
        
        return result
    
    @staticmethod
    def _analyze_text(file_obj, result: Dict, file_ext: str) -> Dict:
        """Analiza archivos de texto (.txt, .md)"""
        try:
            file_obj.seek(0)
            text = file_obj.read().decode('utf-8', errors='ignore')
            
            result['stats']['text_length'] = len(text)
            result['stats']['word_count'] = len(text.split())
            result['stats']['pages'] = max(1, len(text) // 3000)  # ~3000 chars por página
            
            # Validar límite de caracteres
            max_chars = DocumentAnalyzer.MAX_CHARS.get(file_ext)
            if max_chars and len(text) > max_chars:
                result['is_valid'] = False
                result['error'] = f"El archivo tiene {len(text)} caracteres. Máximo permitido: {max_chars} caracteres (~{max_chars//5000} palabras)"
                return result
            
            # Validar que tenga texto útil
            if len(text.strip()) < DocumentAnalyzer.MIN_TEXT_LENGTH:
                result['is_valid'] = False
                result['error'] = "El archivo no contiene suficiente texto útil para procesar"
                return result
            
            logger.info(f"Archivo de texto analizado: {len(text)} caracteres, {len(text.split())} palabras")
            
        except Exception as e:
            result['is_valid'] = False
            result['error'] = f"Error al leer el archivo de texto: {str(e)}"
        
        return result
    
    @staticmethod
    def estimate_tokens(text_length: int) -> int:
        """
        Estima el número de tokens basado en la longitud del texto
        Aproximación: 1 token ≈ 4 caracteres en español
        """
        return text_length // 4
    
    @staticmethod
    def estimate_cost(text_length: int, cost_per_1k_tokens: float = 0.0001) -> float:
        """
        Estima el costo de procesamiento del documento
        
        Args:
            text_length: longitud del texto en caracteres
            cost_per_1k_tokens: costo por 1000 tokens (ajustar según proveedor)
        """
        tokens = DocumentAnalyzer.estimate_tokens(text_length)
        return (tokens / 1000) * cost_per_1k_tokens
